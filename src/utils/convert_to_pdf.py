import tempfile
import os
import subprocess
from pathlib import Path
from docx import Document
import io

def convert_to_pdf(file_bytes: bytes, original_filename: str) -> str:
    """
    Kisi bhi file ko PDF mein convert karta hai aur temporary PDF ka path return karta hai.
    Supported: .pdf, .docx, .txt, .png, .jpg, .jpeg
    """
    ext = Path(original_filename).suffix.lower()
    
    # Agar already PDF hai
    if ext == '.pdf':
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            tmp.write(file_bytes)
            return tmp.name
    
    # DOCX -> PDF (using LibreOffice - best quality)
    elif ext == '.docx':
        from docx import Document
        import io
        from fpdf import FPDF
    
        # Pehle python-docx se text nikalte hain
        try:
            doc = Document(io.BytesIO(file_bytes))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            # Tables se bhi text lo
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
        
            if not full_text:
                raise ValueError("No text found in DOCX")
        
            # Text se PDF banao
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in full_text:
                # Unicode handle
                try:
                    pdf.cell(0, 10, txt=line[:200], ln=True)
                except:
                    safe_line = line.encode('latin-1', 'ignore').decode('latin-1')
                    pdf.cell(0, 10, txt=safe_line[:200], ln=True)
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
            pdf.output(temp_pdf)
            return temp_pdf
        except Exception as e:
            raise ValueError(f"DOCX conversion failed: {str(e)}")
    
    # TXT -> PDF
    elif ext == '.txt':
        from fpdf import FPDF
        text = file_bytes.decode('utf-8', errors='replace')
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in text.split('\n'):
            try:
                pdf.cell(0, 10, txt=line[:200], ln=True)
            except:
                pdf.cell(0, 10, txt=line.encode('latin-1', 'ignore').decode('latin-1'), ln=True)
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        pdf.output(temp_pdf)
        return temp_pdf
    
    # Images -> PDF
    elif ext in ['.png', '.jpg', '.jpeg']:
        import img2pdf
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as img_tmp:
            img_tmp.write(file_bytes)
            img_path = img_tmp.name
        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        with open(temp_pdf, 'wb') as f:
            f.write(img2pdf.convert(img_path, dpi=300))
        os.unlink(img_path)
        return temp_pdf
    
    else:
        raise ValueError(f"Unsupported format: {ext}")