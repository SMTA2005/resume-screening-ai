from docx import Document
import io

def extract_text_from_docx_bytes(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        return "\n".join(full_text)
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""